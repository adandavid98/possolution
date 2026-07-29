import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_design_document():
    doc = Document()
    
    # Page Margins (1 inch everywhere)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Charcoal Dark

    # --- TITLE & ACADEMIC HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = title_p.add_run("UNIVERSIDAD TECNOLÓGICA DE SANTIAGO (UTESA)\n")
    r_univ.font.size = Pt(15)
    r_univ.font.bold = True
    r_univ.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81) # Deep Royal Blue

    r_carrera = title_p.add_run("Carrera de Ingeniería en Sistemas Computacionales\nProyecto Integrador I · INF-225-002\n\n")
    r_carrera.font.size = Pt(12)
    r_carrera.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    r_doc_title = title_p.add_run("DOCUMENTO DE DISEÑO DE SOFTWARE\n")
    r_doc_title.font.size = Pt(22)
    r_doc_title.font.bold = True
    r_doc_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    r_sub = title_p.add_run("Unidad II: Diseño Modular y Diseño de Datos\nSistema de Gestión de Inventario y Ventas para Minimarket La Ruta del Este, S.R.L.\n\n")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # Info Box Table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Asignatura:", "Proyecto Integrador I - INF-225-002"),
        ("Estudiantes:", "Adan Ozoria (1-17-6879), Henderson Branagan (1-18-4481), Sebastian Feliz (1-22-4806)"),
        ("Docente:", "Rossy Almanzar"),
        ("Periodo Académico:", "Cuatrimestre Mayo - Agosto 2026"),
        ("Fecha de Entrega:", "25 de julio de 2026")
    ]
    for row_idx, (label, val) in enumerate(info_data):
        row = info_table.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        p0 = cell_lbl.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        p1 = cell_val.paragraphs[0]
        p1.add_run(val)
        
        set_cell_background(cell_lbl, 'F1F5F9')
        set_cell_background(cell_val, 'FFFFFF')
        set_cell_margins(cell_lbl, 80, 80, 120, 120)
        set_cell_margins(cell_val, 80, 80, 120, 120)
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(4.7)

    doc.add_paragraph("\n")

    # Helper for Headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    # --- PARTE 1: DISEÑO MODULAR ---
    add_heading_1("PARTE 1 — DISEÑO MODULAR")

    add_heading_2("1. Diagrama de Descomposición Modular")
    p_desc = doc.add_paragraph(
        "El subsistema principal de Gestión Operativa y Comercial para Minimarket La Ruta del Este, S.R.L., "
        "se descompone jerárquicamente en módulos independientes especializados. Esta descomposición garantiza "
        "que cada módulo cumpla una responsabilidad única y bien definida dentro del sistema."
    )
    p_desc.paragraph_format.space_after = Pt(10)

    # Modular Hierarchy Box Table
    hierarchy_data = [
        ("Módulo Superior", "0.0 Sistema de Gestión de Inventario y Ventas (Minimarket La Ruta del Este)"),
        ("1.0 Módulo Catálogo de Productos", "1.1 Registro/Edición de Productos\n1.2 Búsqueda SQL Live & Carga Inteligente\n1.3 Clasificación por Deptos/Subdeptos"),
        ("2.0 Módulo Control de Inventario", "2.1 Entradas por Compras a Suplidores\n2.2 Salidas, Pérdidas y Ajustes de Stock\n2.3 Alertas de Stock Mínimo y Reposición"),
        ("3.0 Módulo POS y Ventas Táctiles", "3.1 Carrito de Venta Táctil\n3.2 Checkout Multi-método (Efectivo RD$, Tarjeta, Transferencia, Crédito)\n3.3 Inserción Atómica (OUTPUT INSERTED.id) y Generación Ticket PDF"),
        ("4.0 Módulo Reportes y Analítica", "4.1 Consolidado General e Historial Transaccional\n4.2 Desglose de Ventas por Depto y Subdepto\n4.3 Selector de Fecha In-App y Exportación a PDF/Excel"),
        ("5.0 Módulo Gestión de Caja y Turnos", "5.1 Apertura y Fondo Inicial de Caja\n5.2 Arqueo, Balance y Cierre de Turno")
    ]

    mod_table = doc.add_table(rows=len(hierarchy_data), cols=2)
    mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (m_title, m_funcs) in enumerate(hierarchy_data):
        row = mod_table.rows[r_idx]
        c0, c1 = row.cells[0], row.cells[1]
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(m_title)
        r0.bold = True
        if r_idx == 0:
            r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(c0, '0F4C81')
            set_cell_background(c1, '0F4C81')
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(m_funcs)
            r1.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        else:
            r0.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
            set_cell_background(c0, 'F8FAFC')
            set_cell_background(c1, 'FFFFFF')
            p1 = c1.paragraphs[0]
            p1.add_run(m_funcs)
            
        set_cell_margins(c0, 100, 100, 120, 120)
        set_cell_margins(c1, 100, 100, 120, 120)
        c0.width = Inches(2.4)
        c1.width = Inches(4.1)

    doc.add_paragraph("\n")

    add_heading_2("2. Fichas de Especificación por Módulo")

    fichas = [
        {
            "num": "Ficha 1: Módulo POS y Ventas Táctiles (3.0)",
            "nombre": "Módulo POS y Procesamiento de Ventas (VentaModel / GUI POS)",
            "funcion": "Gestiona la construcción del carrito de compras en tiempo real, efectúa el cálculo del ITBIS (18%), procesa transacciones multi-método (Efectivo RD$, Tarjeta, Transferencia/WhatsApp, Crédito/Fiado), descuenta automáticamente las existencias de inventario, registra ventas atómicamente y emite el comprobante de venta en PDF.",
            "entradas": "ID de caja activa (caja_id), ID de usuario (user_id), lista de productos en carrito (producto_id, cantidad, precio_venta), tipo de pago seleccionado y monto recibido en efectivo.",
            "salidas": "Código de factura único (codigo_factura ej. FAC-20260725131500), comprobante PDF en tickets/, notificación de confirmación en pantalla y registros atómicos en ventas y detalle_ventas.",
            "relaciones": "Se relaciona con Módulo Catálogo (1.0) para obtener precios, Módulo Inventario (2.0) para decrementar stock y Módulo Caja (5.0) para asociar la venta al turno activo."
        },
        {
            "num": "Ficha 2: Módulo Catálogo de Productos (1.0)",
            "nombre": "Módulo de Catálogo y Productos (ProductoModel / GUI Catálogo)",
            "funcion": "Administra el ciclo de vida completo de los artículos comercializados (registro, baja lógica/desactivación, edición de precios, unidades de medida y asignación a departamentos/subdepartamentos). Ofrece búsquedas SQL optimizadas de alta velocidad.",
            "entradas": "Código de barras, nombre, categoría, precio_compra, precio_venta, stock_actual, stock_minimo, unidad_medida, departamento_id, subdepartamento_id y término de búsqueda.",
            "salidas": "Listado estructurado de productos (Top 20 esencial o filtrado dinámico >= 2 caracteres) y confirmación de guardado en la tabla productos.",
            "relaciones": "Suministra información de precios al Módulo POS (3.0), define umbrales de reposición para el Módulo Inventario (2.0) y abastece de jerarquías comerciales al Módulo Reportes (4.0)."
        },
        {
            "num": "Ficha 3: Módulo Control de Inventario (2.0)",
            "nombre": "Módulo de Control y Movimientos de Inventario (MovimientoModel / GUI Inventario)",
            "funcion": "Garantiza la trazabilidad física y lógica de las mercancías registrando entradas por compras, salidas por pérdidas/desperdicios y ajustes manuales. Evalúa continuamente existencias contra el stock mínimo para generar alertas de reposición.",
            "entradas": "producto_id, tipo_movimiento (Entrada, Salida, Ajuste), cantidad, motivo y usuario_id.",
            "salidas": "Historial auditable en la tabla movimientos_inventario, existencias actualizadas en productos y matriz de productos críticos con bajo stock.",
            "relaciones": "Modifica la existencia en Módulo Catálogo (1.0), procesa decrementos automáticos desde Módulo POS (3.0) y abastece de datos al Módulo Reportes (4.0)."
        },
        {
            "num": "Ficha 4: Módulo Reportes y Analítica (4.0)",
            "nombre": "Módulo de Reportes y Generación de Documentos (ReportModel / report_pdf.py / excel_exporter.py)",
            "funcion": "Consolida datos transaccionales en informes analíticos de ventas brutas, ITBIS recaudado, costos y ganancias estimadas. Ofrece filtros por periodo (Día, Semana, Mes, Año, Personalizado) con selector de fecha desplegable e impresiones directas en PDF y Excel.",
            "entradas": "Tipo de reporte (Consolidado General, Ventas por Producto, Ventas por Departamento/Subdepartamento, Cierre de Caja, Movimientos) y rango de fechas ISO YYYY-MM-DD.",
            "salidas": "Dataset analítico en interfaz gráfica, documentos impresos PDF y archivos de hoja de cálculo .xlsx.",
            "relaciones": "Consulta tablas del Módulo POS (3.0), Módulo Catálogo (1.0) y Módulo Gestión de Caja (5.0)."
        }
    ]

    for f in fichas:
        p_fnum = doc.add_paragraph()
        r_fnum = p_fnum.add_run(f["num"])
        r_fnum.bold = True
        r_fnum.font.size = Pt(12)
        r_fnum.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

        f_table = doc.add_table(rows=5, cols=2)
        f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows_info = [
            ("Nombre del Módulo", f["nombre"]),
            ("Función que Cumple", f["funcion"]),
            ("Entradas", f["entradas"]),
            ("Salidas", f["salidas"]),
            ("Relación con Otros Módulos", f["relaciones"])
        ]
        for r_i, (lbl, val) in enumerate(rows_info):
            r_c = f_table.rows[r_i]
            c0, c1 = r_c.cells[0], r_c.cells[1]
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(lbl)
            r0.bold = True
            r0.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
            p1 = c1.paragraphs[0]
            p1.add_run(val)
            
            set_cell_background(c0, 'F1F5F9')
            set_cell_background(c1, 'FFFFFF')
            set_cell_margins(c0, 80, 80, 120, 120)
            set_cell_margins(c1, 80, 80, 120, 120)
            c0.width = Inches(2.0)
            c1.width = Inches(4.5)
            
        doc.add_paragraph()

    add_heading_2("3. Justificación de Cohesión y Acoplamiento")

    cohesiones = [
        ("Módulo POS y Ventas Táctiles (3.0):", "Posee alta cohesión funcional, ya que el 100% de sus componentes y métodos están enfocados de manera exclusiva en la preparación del carrito, cobro multi-método, descuento de inventario y emisión de comprobantes. El acoplamiento se minimizó al interactuar con la base de datos únicamente a través del método atómico VentaModel.procesar_venta(), sin manipular variables directas de otros módulos."),
        ("Módulo Catálogo de Productos (1.0):", "Presenta alta cohesión informática y funcional dedicada a administrar los datos principales de los productos (nombre, precios, departamentos y estados). El acoplamiento con la interfaz de usuario y los módulos secundarios se redujo exponiendo métodos de consulta independientes (obtener_productos()), garantizando que cambios en la UI no impacten la lógica de persistencia."),
        ("Módulo Control de Inventario (2.0):", "Demuestra alta cohesión secuencial y funcional al centralizar todas las reglas de negocio sobre movimientos físicos (entradas, salidas, motivos y umbrales de reposición). Mantiene bajo acoplamiento al recibir solicitudes de descuento como eventos transaccionales numéricos simples (producto_id, cantidad), permitiendo modificar las reglas de almacén sin alterar la UI del POS."),
        ("Módulo Reportes y Analítica (4.0):", "Cuenta con alta cohesión de comunicación, enfocada únicamente en la agregación de métricas transaccionales y su renderizado en formatos legibles (pantalla, PDF y Excel). Garantiza bajo acoplamiento al actuar estrictamente como un módulo de lectura (read-only) sobre la base de datos, asegurando que cualquier cambio en los generadores PDF no interfiera en la ejecución contínua de ventas.")
    ]

    for title, text in cohesiones:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(6)
        r_t = p_c.add_run(title + " ")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
        p_c.add_run(text)

    doc.add_paragraph("\n")

    # --- PARTE 2: DISEÑO DE DATOS ---
    add_heading_1("PARTE 2 — DISEÑO DE DATOS")

    add_heading_2("1. Modelo Entidad-Relación (E-R)")
    doc.add_paragraph(
        "El modelo relacional del sistema está compuesto por 6 entidades principales diseñadas bajo norma "
        "de tercera forma normal (3FN), garantizando integridad referencial y soporte de transacciones atómicas."
    )

    er_entities = [
        ("PRODUCTOS", "id (PK), codigo_barras (UK), nombre, precio_compra, precio_venta, stock_actual, stock_minimo, unidad_medida, estado, departamento_id (FK), subdepartamento_id (FK)"),
        ("VENTAS", "id (PK - OUTPUT INSERTED.id), codigo_factura (UK), fecha_venta, caja_id (FK), usuario_id (FK), cliente_nombre, tipo_pago, subtotal, itbis, total"),
        ("DETALLE_VENTAS", "id (PK), venta_id (FK), producto_id (FK), cantidad, precio_unitario, subtotal"),
        ("MOVIMIENTOS_INVENTARIO", "id (PK), producto_id (FK), tipo_movimiento, cantidad, motivo, fecha_movimiento, usuario_id (FK)"),
        ("CAJAS", "id (PK), nombre_caja, estado, fondo_inicial"),
        ("TURNOS_CAJA", "id (PK), caja_id (FK), usuario_id (FK), fecha_apertura, fecha_cierre, monto_apertura, monto_cierre, estado")
    ]

    er_table = doc.add_table(rows=len(er_entities)+1, cols=2)
    er_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    hdr = er_table.rows[0]
    set_cell_background(hdr.cells[0], '0F4C81')
    set_cell_background(hdr.cells[1], '0F4C81')
    p_h0 = hdr.cells[0].paragraphs[0].add_run("Entidad")
    p_h0.bold = True
    p_h0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_h1 = hdr.cells[1].paragraphs[0].add_run("Atributos y Claves (PK / FK / UK)")
    p_h1.bold = True
    p_h1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (ent, attrs) in enumerate(er_entities, start=1):
        row = er_table.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        
        p0 = c0.paragraphs[0].add_run(ent)
        p0.bold = True
        p0.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
        
        c1.paragraphs[0].add_run(attrs)
        
        set_cell_background(c0, 'F8FAFC' if idx % 2 == 1 else 'FFFFFF')
        set_cell_background(c1, 'F8FAFC' if idx % 2 == 1 else 'FFFFFF')
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)

    doc.add_paragraph("\n")
    p_card = doc.add_paragraph()
    p_card.add_run("Relaciones y Cardinalidades Centrales:\n").bold = True
    p_card.add_run("• PRODUCTOS (1) ─── (N) DETALLE_VENTAS : Un producto puede ser vendido en múltiples detalles de venta.\n")
    p_card.add_run("• VENTAS (1) ─── (N) DETALLE_VENTAS : Una venta contiene de 1 a N líneas de detalle de producto.\n")
    p_card.add_run("• PRODUCTOS (1) ─── (N) MOVIMIENTOS_INVENTARIO : Un producto registra múltiples entradas, salidas y ajustes.\n")
    p_card.add_run("• CAJAS (1) ─── (N) VENTAS : Una caja registra N transacciones de venta asociadas.\n")

    add_heading_2("2. Diccionario de Datos")

    dict_tables = [
        {
            "name": "Tabla 1: Entidad 'productos' (Catálogo de Artículos)",
            "rows": [
                ("id", "INT", "4 bytes", "PK, IDENTITY(1,1)", "Identificador único del producto."),
                ("codigo_barras", "VARCHAR", "50", "UNIQUE, NOT NULL", "Código EAN/UPC escaneable."),
                ("nombre", "VARCHAR", "150", "NOT NULL", "Descripción comercial del producto."),
                ("precio_compra", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Costo de adquisición con suplidor."),
                ("precio_venta", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Precio final al consumidor."),
                ("stock_actual", "INT", "4 bytes", "NOT NULL, DEFAULT 0", "Existencia disponible en inventario."),
                ("stock_minimo", "INT", "4 bytes", "NOT NULL, DEFAULT 5", "Umbral crítico para alerta de reposición."),
                ("unidad_medida", "VARCHAR", "20", "NOT NULL, DEFAULT 'Unidad'", "Presentación (Unidad, Libra, Botella)."),
                ("estado", "VARCHAR", "20", "NOT NULL, DEFAULT 'Activo'", "Estado operacional (Activo/Inactivo)."),
                ("departamento_id", "INT", "4 bytes", "FK (departamentos.id)", "Enlace al departamento comercial."),
                ("subdepartamento_id", "INT", "4 bytes", "FK (subdepartamentos.id)", "Enlace al subdepartamento comercial.")
            ]
        },
        {
            "name": "Tabla 2: Entidad 'ventas' (Encabezado de Transacción)",
            "rows": [
                ("id", "INT", "4 bytes", "PK, IDENTITY(1,1)", "Clave atómica devuelta por OUTPUT INSERTED.id."),
                ("codigo_factura", "VARCHAR", "50", "UNIQUE, NOT NULL", "Código de comprobante (FAC-YYYYMMDDHHMMSS)."),
                ("fecha_venta", "DATETIME", "8 bytes", "NOT NULL, DEFAULT GETDATE()", "Fecha y hora de confirmación del cobro."),
                ("caja_id", "INT", "4 bytes", "FK (cajas.id), NOT NULL", "Caja donde se procesó la venta."),
                ("usuario_id", "INT", "4 bytes", "FK (usuarios.id), NOT NULL", "Cajero/usuario que realizó el cobro."),
                ("cliente_nombre", "VARCHAR", "100", "NOT NULL, DEFAULT 'Cliente General'", "Nombre o RNC del cliente."),
                ("tipo_pago", "VARCHAR", "30", "NOT NULL", "Efectivo, Tarjeta, Transferencia, Credito."),
                ("subtotal", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Suma imponible de los artículos."),
                ("itbis", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Monto del impuesto ITBIS (18%)."),
                ("total", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Monto total bruto cobrado.")
            ]
        },
        {
            "name": "Tabla 3: Entidad 'detalle_ventas' (Líneas de Producto Vendido)",
            "rows": [
                ("id", "INT", "4 bytes", "PK, IDENTITY(1,1)", "Identificador único de la línea de detalle."),
                ("venta_id", "INT", "4 bytes", "FK (ventas.id), NOT NULL", "Enlace a la venta padre correspondiente."),
                ("producto_id", "INT", "4 bytes", "FK (productos.id), NOT NULL", "Enlace al producto comercializado."),
                ("cantidad", "INT", "4 bytes", "NOT NULL, > 0", "Unidades vendidas del artículo."),
                ("precio_unitario", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Precio unitario aplicado en la transacción."),
                ("subtotal", "DECIMAL", "(10,2)", "NOT NULL, >= 0", "Importe de la línea (cantidad * precio).")
            ]
        }
    ]

    for dt in dict_tables:
        p_dt = doc.add_paragraph()
        r_dt = p_dt.add_run(dt["name"])
        r_dt.bold = True
        r_dt.font.size = Pt(11)
        r_dt.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

        tbl = doc.add_table(rows=len(dt["rows"])+1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Table Header
        h_row = tbl.rows[0]
        h_headers = ["Campo", "Tipo Dato", "Tamaño", "Restricción", "Descripción"]
        col_w = [Inches(1.3), Inches(1.0), Inches(0.8), Inches(1.6), Inches(1.8)]
        
        for c_idx, text in enumerate(h_headers):
            cell = h_row.cells[c_idx]
            set_cell_background(cell, '0F4C81')
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)
            cell.width = col_w[c_idx]

        for r_idx, (col_name, data_type, size_val, restr, desc) in enumerate(dt["rows"], start=1):
            row = tbl.rows[r_idx]
            bg_hex = 'F8FAFC' if r_idx % 2 == 1 else 'FFFFFF'
            
            vals = [col_name, data_type, size_val, restr, desc]
            for c_idx, val_str in enumerate(vals):
                cell = row.cells[c_idx]
                set_cell_background(cell, bg_hex)
                set_cell_margins(cell, 60, 60, 80, 80)
                p = cell.paragraphs[0]
                r = p.add_run(val_str)
                r.font.size = Pt(9.5)
                if c_idx == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
                cell.width = col_w[c_idx]

        doc.add_paragraph()

    # Save outputs
    output_filename = "Diseno_de_Software_Minimarket_La_Ruta_del_Este.docx"
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")
    
    # Also save to Artifacts folder
    artifact_dir = r"C:\Users\Adan\.gemini\antigravity\brain\b9fa430f-b95a-4f53-a66e-abcf647967d2"
    if os.path.exists(artifact_dir):
        artifact_path = os.path.join(artifact_dir, output_filename)
        doc.save(artifact_path)
        print(f"Artifact document saved to: {artifact_path}")

if __name__ == "__main__":
    create_design_document()
